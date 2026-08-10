"""File guide.

- Use: Contains integration tests for document processing integration behavior.
- Usage: Run this file with pytest when checking related behavior.
- Duties: Builds test data, calls the public API, and checks expected results.
- Depends on: Project modules: openagentlab.skills,
  openagentlab.skills.document_processing,
  openagentlab.skills.document_processing.tools.csv_reader,
  openagentlab.skills.document_processing.tools.docx_reader,
  openagentlab.skills.document_processing.tools.excel_sheet_reader, and 6 more.
"""

import json
from datetime import date, datetime
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook
from pydantic import BaseModel, ValidationError

from openagentlab.skills import SkillRegistry
from openagentlab.skills.document_processing import DocumentProcessingSkill
from openagentlab.skills.document_processing.tools.csv_reader import (
    CSVReaderError,
    CSVReaderInput,
    CSVReaderOutput,
)
from openagentlab.skills.document_processing.tools.docx_reader import (
    DOCXReaderError,
    DOCXReaderInput,
    DOCXReaderOutput,
)
from openagentlab.skills.document_processing.tools.excel_sheet_reader import (
    ExcelSheetReaderError,
    ExcelSheetReaderInput,
    ExcelSheetReaderOutput,
)
from openagentlab.skills.document_processing.tools.excel_workbook_reader import (
    ExcelWorkbookReaderError,
    ExcelWorkbookReaderInput,
    ExcelWorkbookReaderOutput,
)
from openagentlab.skills.document_processing.tools.json_reader import (
    JSONReaderError,
    JSONReaderInput,
    JSONReaderOutput,
)
from openagentlab.skills.document_processing.tools.pdf_reader import (
    PDFReaderError,
    PDFReaderInput,
    PDFReaderOutput,
)
from openagentlab.skills.document_processing.tools.text_reader import (
    TextReaderError,
    TextReaderInput,
    TextReaderOutput,
)
from openagentlab.skills.registry import DuplicateSkillError
from openagentlab.skills.tool import BaseTool

EXPECTED_DECLARED_CAPABILITIES = (
    "document.read.pdf",
    "document.read.excel",
    "document.read.excel.workbook",
    "document.read.excel.sheet",
    "document.read.csv",
    "document.read.text",
    "document.read.json",
    "document.read.docx",
)
EXPECTED_EXECUTABLE_CAPABILITIES = (
    "document.read.pdf",
    "document.read.excel.workbook",
    "document.read.excel.sheet",
    "document.read.csv",
    "document.read.text",
    "document.read.json",
    "document.read.docx",
)
EXPECTED_TOOL_NAMES = (
    "pdf_reader",
    "excel_workbook_reader",
    "excel_sheet_reader",
    "csv_reader",
    "text_reader",
    "json_reader",
    "docx_reader",
)


def get_required_tool(skill: DocumentProcessingSkill, name: str) -> BaseTool:
    tool = skill.get_tool(name)
    assert tool is not None
    return tool


def write_text_pdf(
    path: Path,
    page_texts: tuple[str, ...],
    metadata: dict[str, str] | None = None,
) -> None:
    objects: list[bytes] = []
    page_numbers: list[int] = []

    def add_object(content: bytes) -> int:
        objects.append(content)
        return len(objects)

    add_object(b"<< /Type /Catalog /Pages 2 0 R >>")
    add_object(b"")
    add_object(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    for text in page_texts:
        page_number = len(objects) + 1
        content_number = len(objects) + 2
        page_numbers.append(page_number)

        add_object(
            (
                f"<< /Type /Page /Parent 2 0 R "
                f"/MediaBox [0 0 612 792] "
                f"/Resources << /Font << /F1 3 0 R >> >> "
                f"/Contents {content_number} 0 R >>"
            ).encode()
        )

        content = f"BT\n/F1 12 Tf\n72 720 Td\n{pdf_string(text)} Tj\nET\n".encode()
        add_object(
            b"<< /Length "
            + str(len(content)).encode()
            + b" >>\nstream\n"
            + content
            + b"endstream"
        )

    kids = " ".join(f"{page_number} 0 R" for page_number in page_numbers)
    objects[1] = (
        f"<< /Type /Pages /Kids [{kids}] /Count {len(page_numbers)} >>".encode()
    )

    info_number = None
    if metadata:
        entries = " ".join(
            f"/{key} {pdf_string(value)}" for key, value in metadata.items()
        )
        info_number = add_object(f"<< {entries} >>".encode())

    write_pdf_objects(path, objects, info_number)


def pdf_string(value: str) -> str:
    escaped = value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    return f"({escaped})"


def write_pdf_objects(
    path: Path,
    objects: list[bytes],
    info_number: int | None,
) -> None:
    content = b"%PDF-1.4\n"
    offsets = []

    for object_number, object_content in enumerate(objects, start=1):
        offsets.append(len(content))
        content += f"{object_number} 0 obj\n".encode() + object_content + b"\nendobj\n"

    xref_start = len(content)
    content += f"xref\n0 {len(objects) + 1}\n".encode()
    content += b"0000000000 65535 f \n"

    for offset in offsets:
        content += f"{offset:010d} 00000 n \n".encode()

    trailer = f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R"
    if info_number is not None:
        trailer += f" /Info {info_number} 0 R"
    trailer += f" >>\nstartxref\n{xref_start}\n%%EOF\n"

    path.write_bytes(content + trailer.encode())


def write_workbook(path: Path) -> None:
    workbook = Workbook()

    summary = workbook.active
    summary.title = "Summary"
    summary.append(["name", "amount", "active", "when"])
    summary.append(["alpha", 10, True, datetime(2026, 8, 8, 9, 30, 0)])
    summary.append(["beta", 12.5, False, date(2026, 8, 9)])
    summary.append(["empty-middle", None, True, "tail"])

    formula = workbook.create_sheet("Formula")
    formula["A1"] = 1
    formula["A2"] = 2
    formula["A3"] = "=SUM(A1:A2)"

    sparse = workbook.create_sheet("Sparse")
    sparse["C4"] = "corner"

    workbook.create_sheet("Empty")

    workbook.properties.title = "Integration Workbook"
    workbook.properties.subject = "Document Processing"
    workbook.properties.creator = "OpenAgentLab"
    workbook.properties.description = "Integration fixture"
    workbook.properties.created = datetime(2026, 8, 8, 9, 30, 0)

    workbook.save(path)
    workbook.close()


def write_csv(path: Path, content: str, encoding: str = "utf-8") -> None:
    path.write_text(content, encoding=encoding, newline="")


def write_plain_text(path: Path, content: str, encoding: str = "utf-8") -> None:
    path.write_text(content, encoding=encoding)


def write_json_value(path: Path, value: object, encoding: str = "utf-8") -> None:
    path.write_text(json.dumps(value, ensure_ascii=False), encoding=encoding)


def write_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Integration paragraph")
    document.add_paragraph("  Second paragraph  ")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "name"
    table.cell(0, 1).text = "value"
    table.cell(1, 0).text = "alpha"
    table.cell(1, 1).text = "42"

    document.core_properties.title = "Integration DOCX"
    document.core_properties.author = "OpenAgentLab"
    document.core_properties.created = datetime(2026, 8, 8, 9, 30, 0)

    document.save(path)


def test_document_processing_skill_contract() -> None:
    skill = DocumentProcessingSkill()

    assert skill.name == "document_processing"
    assert skill.metadata.name == "document_processing"
    assert skill.metadata.description
    assert skill.metadata.version
    assert skill.instructions
    assert skill.capabilities == EXPECTED_DECLARED_CAPABILITIES
    assert skill.executable_capabilities == EXPECTED_EXECUTABLE_CAPABILITIES
    assert skill.dependencies == ()
    assert tuple(tool.name for tool in skill.tools) == EXPECTED_TOOL_NAMES
    assert skill.get_tool("unknown_tool") is None

    tool_names = [tool.name for tool in skill.tools]
    tool_capabilities = [tool.capability for tool in skill.tools]

    assert len(tool_names) == len(set(tool_names))
    assert len(tool_capabilities) == len(set(tool_capabilities))
    assert len(tool_capabilities) == len(skill.executable_capabilities)

    for capability in skill.executable_capabilities:
        matching_tools = [tool for tool in skill.tools if tool.capability == capability]
        assert len(matching_tools) == 1

    for tool in skill.tools:
        assert isinstance(tool, BaseTool)
        assert tool.name
        assert tool.description
        assert tool.capability
        assert tool.capability in skill.capabilities


def test_document_processing_skill_registry_contract() -> None:
    registry = SkillRegistry()
    skill = DocumentProcessingSkill()

    registry.register(skill)

    assert registry.get("document_processing") is skill
    assert registry.get("unknown") is None
    assert registry.list_skills() == (skill,)
    assert registry.find_by_capability("document.read.excel") == (skill,)
    assert registry.find_by_capability("document.read.csv") == (skill,)
    assert registry.find_by_capability("unknown.capability") == ()

    with pytest.raises(DuplicateSkillError, match="document_processing"):
        registry.register(DocumentProcessingSkill())


def test_pdf_reader_executes_through_skill(tmp_path: Path) -> None:
    path = tmp_path / "sample.pdf"
    write_text_pdf(
        path,
        ("First page", "Second page", ""),
        {"Title": "Integration PDF", "Author": "OpenAgentLab"},
    )
    tool = get_required_tool(DocumentProcessingSkill(), "pdf_reader")

    result = tool.execute(PDFReaderInput(path=path))

    assert isinstance(result, PDFReaderOutput)
    assert result.path == path
    assert result.page_count == 3
    assert [page.page_number for page in result.pages] == [1, 2, 3]
    assert "First page" in result.pages[0].text
    assert "Second page" in result.pages[1].text
    assert result.pages[2].text == ""
    assert result.metadata["Title"] == "Integration PDF"
    assert result.metadata["Author"] == "OpenAgentLab"


@pytest.mark.parametrize(
    ("filename", "content", "input_filename", "error_type"),
    [
        ("missing.pdf", None, "missing.pdf", FileNotFoundError),
        ("not-pdf.txt", b"plain text", "not-pdf.txt", ValueError),
        ("broken.pdf", b"not a pdf", "broken.pdf", PDFReaderError),
    ],
)
def test_pdf_reader_failures_propagate_through_skill(
    tmp_path: Path,
    filename: str,
    content: bytes | None,
    input_filename: str,
    error_type: type[Exception],
) -> None:
    if content is not None:
        (tmp_path / filename).write_bytes(content)
    tool = get_required_tool(DocumentProcessingSkill(), "pdf_reader")

    with pytest.raises(error_type):
        tool.execute(PDFReaderInput(path=tmp_path / input_filename))


def test_excel_workbook_reader_executes_through_skill(tmp_path: Path) -> None:
    path = tmp_path / "workbook.xlsx"
    write_workbook(path)
    tool = get_required_tool(DocumentProcessingSkill(), "excel_workbook_reader")

    result = tool.execute(ExcelWorkbookReaderInput(path=path))

    assert isinstance(result, ExcelWorkbookReaderOutput)
    assert result.path == path
    assert result.sheet_count == 4
    assert [(sheet.index, sheet.name) for sheet in result.sheets] == [
        (1, "Summary"),
        (2, "Formula"),
        (3, "Sparse"),
        (4, "Empty"),
    ]
    assert result.sheets[0].max_row == 4
    assert result.sheets[0].max_column == 4
    assert result.sheets[2].max_row == 4
    assert result.sheets[2].max_column == 3
    assert result.sheets[3].max_row == 1
    assert result.sheets[3].max_column == 1
    assert result.metadata["title"] == "Integration Workbook"
    assert result.metadata["creator"] == "OpenAgentLab"
    assert "worksheets" not in result.model_dump()


@pytest.mark.parametrize(
    ("filename", "content", "input_filename", "error_type"),
    [
        ("missing.xlsx", None, "missing.xlsx", FileNotFoundError),
        ("workbook.xls", b"not supported", "workbook.xls", ValueError),
        ("broken.xlsx", b"not a workbook", "broken.xlsx", ExcelWorkbookReaderError),
    ],
)
def test_excel_workbook_reader_failures_propagate_through_skill(
    tmp_path: Path,
    filename: str,
    content: bytes | None,
    input_filename: str,
    error_type: type[Exception],
) -> None:
    if content is not None:
        (tmp_path / filename).write_bytes(content)
    tool = get_required_tool(DocumentProcessingSkill(), "excel_workbook_reader")

    with pytest.raises(error_type):
        tool.execute(ExcelWorkbookReaderInput(path=tmp_path / input_filename))


def test_excel_sheet_reader_executes_through_skill(tmp_path: Path) -> None:
    path = tmp_path / "sheets.xlsx"
    write_workbook(path)
    tool = get_required_tool(DocumentProcessingSkill(), "excel_sheet_reader")

    result = tool.execute(ExcelSheetReaderInput(path=path, sheet_name="Summary"))

    assert isinstance(result, ExcelSheetReaderOutput)
    assert result.path == path
    assert result.sheet_name == "Summary"
    assert result.row_count == 4
    assert result.column_count == 4
    assert result.truncated is False
    assert result.rows[0] == ("name", "amount", "active", "when")
    assert result.rows[1] == ("alpha", 10, True, "2026-08-08T09:30:00")
    assert result.rows[2] == ("beta", 12.5, False, "2026-08-09T00:00:00")
    assert result.rows[3] == ("empty-middle", None, True, "tail")


def test_excel_sheet_reader_preserves_formulas_and_selects_exact_sheet(
    tmp_path: Path,
) -> None:
    path = tmp_path / "formula.xlsx"
    write_workbook(path)
    tool = get_required_tool(DocumentProcessingSkill(), "excel_sheet_reader")

    result = tool.execute(ExcelSheetReaderInput(path=path, sheet_name="Formula"))

    assert result.rows == ((1,), (2,), ("=SUM(A1:A2)",))
    assert result.rows[-1][0] != 3

    with pytest.raises(ExcelSheetReaderError, match="Worksheet does not exist"):
        tool.execute(ExcelSheetReaderInput(path=path, sheet_name="formula"))


def test_excel_sheet_reader_max_rows_and_sparse_shape_through_skill(
    tmp_path: Path,
) -> None:
    path = tmp_path / "limited.xlsx"
    write_workbook(path)
    tool = get_required_tool(DocumentProcessingSkill(), "excel_sheet_reader")

    limited = tool.execute(
        ExcelSheetReaderInput(path=path, sheet_name="Summary", max_rows=2)
    )
    sparse = tool.execute(ExcelSheetReaderInput(path=path, sheet_name="Sparse"))

    assert limited.row_count == 4
    assert limited.column_count == 4
    assert limited.rows == (
        ("name", "amount", "active", "when"),
        ("alpha", 10, True, "2026-08-08T09:30:00"),
    )
    assert limited.truncated is True
    assert sparse.row_count == 4
    assert sparse.column_count == 3
    assert sparse.rows[-1] == (None, None, "corner")


def test_excel_sheet_reader_failures_propagate_through_skill(tmp_path: Path) -> None:
    workbook_path = tmp_path / "workbook.xlsx"
    broken_path = tmp_path / "broken.xlsx"
    write_workbook(workbook_path)
    broken_path.write_bytes(b"not a workbook")
    tool = get_required_tool(DocumentProcessingSkill(), "excel_sheet_reader")

    with pytest.raises(FileNotFoundError):
        tool.execute(
            ExcelSheetReaderInput(path=tmp_path / "missing.xlsx", sheet_name="Summary")
        )

    with pytest.raises(ExcelSheetReaderError, match="Worksheet does not exist"):
        tool.execute(ExcelSheetReaderInput(path=workbook_path, sheet_name="Missing"))

    with pytest.raises(ExcelSheetReaderError, match="Could not read Excel workbook"):
        tool.execute(ExcelSheetReaderInput(path=broken_path, sheet_name="Summary"))

    with pytest.raises(ValidationError):
        ExcelSheetReaderInput(path=workbook_path, sheet_name="Summary", max_rows=0)


@pytest.mark.parametrize(
    ("delimiter", "content"),
    [
        (",", "a,b,c\n1,2,3\n"),
        (";", "a;b;c\n1;2;3\n"),
        ("\t", "a\tb\tc\n1\t2\t3\n"),
        ("|", "a|b|c\n1|2|3\n"),
    ],
)
def test_csv_reader_reads_delimited_files_through_skill(
    tmp_path: Path,
    delimiter: str,
    content: str,
) -> None:
    path = tmp_path / "data.csv"
    write_csv(path, content)
    tool = get_required_tool(DocumentProcessingSkill(), "csv_reader")

    result = tool.execute(CSVReaderInput(path=path, delimiter=delimiter))

    assert isinstance(result, CSVReaderOutput)
    assert result.rows == (("a", "b", "c"), ("1", "2", "3"))
    assert result.row_count == 2
    assert result.column_count == 3
    assert result.delimiter == delimiter


def test_csv_reader_preserves_textual_values_and_structure_through_skill(
    tmp_path: Path,
) -> None:
    path = tmp_path / "values.csv"
    write_csv(
        path,
        (
            "name,age,active,date,note,empty\n"
            'Alice,42,true,2026-08-08,"hello, friend",\n'
            'Bob,7,false,2026-08-09,"quote ""inside""",tail\n'
            "short,row\n"
        ),
    )
    tool = get_required_tool(DocumentProcessingSkill(), "csv_reader")

    result = tool.execute(CSVReaderInput(path=path, delimiter=","))

    assert result.rows[0] == ("name", "age", "active", "date", "note", "empty")
    assert result.rows[1] == (
        "Alice",
        "42",
        "true",
        "2026-08-08",
        "hello, friend",
        "",
    )
    assert result.rows[2] == (
        "Bob",
        "7",
        "false",
        "2026-08-09",
        'quote "inside"',
        "tail",
    )
    assert result.rows[3] == ("short", "row")
    assert result.column_count == 6
    assert isinstance(result.rows[1][1], str)
    assert isinstance(result.rows[1][2], str)
    assert isinstance(result.rows[1][3], str)


def test_csv_reader_detection_encoding_and_max_rows_through_skill(
    tmp_path: Path,
) -> None:
    path = tmp_path / "detected.csv"
    latin_path = tmp_path / "latin.csv"
    write_csv(path, "a|b\n1|2\n3|4\n")
    write_csv(latin_path, "name\nAndré\n", encoding="latin-1")
    tool = get_required_tool(DocumentProcessingSkill(), "csv_reader")

    detected = tool.execute(CSVReaderInput(path=path, max_rows=2))
    latin = tool.execute(
        CSVReaderInput(path=latin_path, delimiter=",", encoding="latin-1")
    )

    assert detected.delimiter == "|"
    assert detected.row_count == 3
    assert detected.rows == (("a", "b"), ("1", "2"))
    assert detected.truncated is True
    assert latin.encoding == "latin-1"
    assert latin.rows == (("name",), ("André",))


def test_csv_reader_failures_propagate_through_skill(tmp_path: Path) -> None:
    tool = get_required_tool(DocumentProcessingSkill(), "csv_reader")
    unsupported = tmp_path / "data.txt"
    unknown_delimiter = tmp_path / "unknown.csv"
    bad_encoding = tmp_path / "bad-encoding.csv"
    malformed = tmp_path / "malformed.csv"
    unsupported.write_text("a,b\n", encoding="utf-8")
    unknown_delimiter.write_text("single column\nanother row\n", encoding="utf-8")
    bad_encoding.write_bytes("name\nAndré\n".encode("latin-1"))
    malformed.write_text('name,note\nAlice,"unterminated\n', encoding="utf-8")

    with pytest.raises(FileNotFoundError):
        tool.execute(CSVReaderInput(path=tmp_path / "missing.csv", delimiter=","))

    with pytest.raises(IsADirectoryError):
        tool.execute(CSVReaderInput(path=tmp_path, delimiter=","))

    with pytest.raises(ValueError, match=".csv extension"):
        tool.execute(CSVReaderInput(path=unsupported, delimiter=","))

    with pytest.raises(ValidationError):
        CSVReaderInput(path=unsupported, delimiter="::")

    with pytest.raises(CSVReaderError, match="Could not detect CSV delimiter"):
        tool.execute(CSVReaderInput(path=unknown_delimiter))

    with pytest.raises(CSVReaderError, match="Unknown CSV encoding"):
        tool.execute(CSVReaderInput(path=unknown_delimiter, encoding="not-a-codec"))

    with pytest.raises(CSVReaderError, match="could not be decoded"):
        tool.execute(CSVReaderInput(path=bad_encoding, delimiter=","))

    with pytest.raises(CSVReaderError, match="Could not read CSV file"):
        tool.execute(CSVReaderInput(path=malformed, delimiter=","))


def test_text_reader_executes_through_skill(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    content = "  hello\n\nworld café\ntrailing   "
    write_plain_text(path, content)
    tool = get_required_tool(DocumentProcessingSkill(), "text_reader")

    result = tool.execute(TextReaderInput(path=path, max_chars=10))

    assert isinstance(result, TextReaderOutput)
    assert result.path == path
    assert result.text == content[:10]
    assert result.encoding == "utf-8"
    assert result.char_count == len(content)
    assert result.line_count == 4
    assert result.truncated is True


def test_text_reader_failures_propagate_through_skill(tmp_path: Path) -> None:
    tool = get_required_tool(DocumentProcessingSkill(), "text_reader")
    unsupported = tmp_path / "notes.md"
    bad_encoding = tmp_path / "bad-encoding.txt"
    unsupported.write_text("# markdown", encoding="utf-8")
    bad_encoding.write_bytes("André".encode("latin-1"))

    with pytest.raises(FileNotFoundError):
        tool.execute(TextReaderInput(path=tmp_path / "missing.txt"))

    with pytest.raises(IsADirectoryError):
        tool.execute(TextReaderInput(path=tmp_path))

    with pytest.raises(ValueError, match=".txt extension"):
        tool.execute(TextReaderInput(path=unsupported))

    with pytest.raises(TextReaderError, match="Unknown text encoding"):
        tool.execute(TextReaderInput(path=bad_encoding, encoding="not-a-codec"))

    with pytest.raises(TextReaderError, match="could not be decoded"):
        tool.execute(TextReaderInput(path=bad_encoding))

    with pytest.raises(ValidationError):
        TextReaderInput(path=bad_encoding, max_chars=0)


def test_json_reader_executes_through_skill(tmp_path: Path) -> None:
    path = tmp_path / "document.json"
    value = {
        "name": "OpenAgentLab",
        "count": 42,
        "active": True,
        "date": "2026-08-10",
        "items": [1, "2", False, None, {"nested": "yes"}],
    }
    write_json_value(path, value)
    tool = get_required_tool(DocumentProcessingSkill(), "json_reader")

    result = tool.execute(JSONReaderInput(path=path))

    assert isinstance(result, JSONReaderOutput)
    assert result.path == path
    assert result.data == value
    assert result.root_type == "object"
    assert result.item_count == 5
    assert result.encoding == "utf-8"
    assert isinstance(result.data["count"], int)
    assert isinstance(result.data["date"], str)


def test_json_reader_failures_propagate_through_skill(tmp_path: Path) -> None:
    tool = get_required_tool(DocumentProcessingSkill(), "json_reader")
    malformed = tmp_path / "malformed.json"
    unsupported = tmp_path / "document.jsonl"
    bad_encoding = tmp_path / "bad-encoding.json"
    malformed.write_text('{"missing": ', encoding="utf-8")
    unsupported.write_text('{"ok": true}\n', encoding="utf-8")
    bad_encoding.write_bytes('{"name": "André"}'.encode("latin-1"))

    with pytest.raises(FileNotFoundError):
        tool.execute(JSONReaderInput(path=tmp_path / "missing.json"))

    with pytest.raises(IsADirectoryError):
        tool.execute(JSONReaderInput(path=tmp_path))

    with pytest.raises(ValueError, match=".json extension"):
        tool.execute(JSONReaderInput(path=unsupported))

    with pytest.raises(JSONReaderError, match="Malformed JSON file"):
        tool.execute(JSONReaderInput(path=malformed))

    with pytest.raises(JSONReaderError, match="Unknown JSON encoding"):
        tool.execute(JSONReaderInput(path=malformed, encoding="not-a-codec"))

    with pytest.raises(JSONReaderError, match="could not be decoded"):
        tool.execute(JSONReaderInput(path=bad_encoding))


def test_docx_reader_executes_through_skill(tmp_path: Path) -> None:
    path = tmp_path / "document.docx"
    write_docx(path)
    tool = get_required_tool(DocumentProcessingSkill(), "docx_reader")

    result = tool.execute(DOCXReaderInput(path=path))

    assert isinstance(result, DOCXReaderOutput)
    assert result.path == path
    assert result.paragraph_count == 2
    assert result.table_count == 1
    assert [paragraph.index for paragraph in result.paragraphs] == [1, 2]
    assert result.paragraphs[0].text == "Integration paragraph"
    assert result.paragraphs[1].text == "  Second paragraph  "
    assert result.tables[0].index == 1
    assert result.tables[0].rows == (("name", "value"), ("alpha", "42"))
    assert isinstance(result.tables[0].rows[1][1], str)
    assert result.metadata["title"] == "Integration DOCX"
    assert result.metadata["author"] == "OpenAgentLab"


def test_docx_reader_failures_propagate_through_skill(tmp_path: Path) -> None:
    tool = get_required_tool(DocumentProcessingSkill(), "docx_reader")
    unsupported = tmp_path / "document.doc"
    broken = tmp_path / "broken.docx"
    unsupported.write_bytes(b"not supported")
    broken.write_bytes(b"not a real docx")

    with pytest.raises(FileNotFoundError):
        tool.execute(DOCXReaderInput(path=tmp_path / "missing.docx"))

    with pytest.raises(IsADirectoryError):
        tool.execute(DOCXReaderInput(path=tmp_path))

    with pytest.raises(ValueError, match=".docx extension"):
        tool.execute(DOCXReaderInput(path=unsupported))

    with pytest.raises(DOCXReaderError, match="Could not read DOCX file"):
        tool.execute(DOCXReaderInput(path=broken))


def test_representative_outputs_serialize_without_library_objects(
    tmp_path: Path,
) -> None:
    pdf_path = tmp_path / "document.pdf"
    workbook_path = tmp_path / "workbook.xlsx"
    csv_path = tmp_path / "data.csv"
    text_path = tmp_path / "notes.txt"
    json_path = tmp_path / "document.json"
    docx_path = tmp_path / "document.docx"
    write_text_pdf(pdf_path, ("Serializable PDF",), {"Title": "Serializable"})
    write_workbook(workbook_path)
    write_csv(csv_path, "a,b\n1,2\n")
    write_plain_text(text_path, "serializable text")
    write_json_value(json_path, {"serializable": [1, True, None]})
    write_docx(docx_path)
    skill = DocumentProcessingSkill()

    outputs = (
        get_required_tool(skill, "pdf_reader").execute(PDFReaderInput(path=pdf_path)),
        get_required_tool(skill, "excel_workbook_reader").execute(
            ExcelWorkbookReaderInput(path=workbook_path)
        ),
        get_required_tool(skill, "excel_sheet_reader").execute(
            ExcelSheetReaderInput(path=workbook_path, sheet_name="Summary", max_rows=2)
        ),
        get_required_tool(skill, "csv_reader").execute(
            CSVReaderInput(path=csv_path, delimiter=",")
        ),
        get_required_tool(skill, "text_reader").execute(
            TextReaderInput(path=text_path)
        ),
        get_required_tool(skill, "json_reader").execute(
            JSONReaderInput(path=json_path)
        ),
        get_required_tool(skill, "docx_reader").execute(
            DOCXReaderInput(path=docx_path)
        ),
    )

    for output in outputs:
        assert isinstance(output, BaseModel)
        dumped = output.model_dump()
        encoded = json.loads(output.model_dump_json())
        assert dumped
        assert isinstance(encoded["path"], str)
        assert "openpyxl" not in str(encoded)
        assert "pypdf" not in str(encoded)
        assert "docx." not in str(encoded)


def test_repeated_execution_is_deterministic_through_skill(tmp_path: Path) -> None:
    pdf_path = tmp_path / "document.pdf"
    workbook_path = tmp_path / "workbook.xlsx"
    csv_path = tmp_path / "data.csv"
    text_path = tmp_path / "notes.txt"
    json_path = tmp_path / "document.json"
    docx_path = tmp_path / "document.docx"
    write_text_pdf(pdf_path, ("Repeatable PDF",), {"Title": "Repeatable"})
    write_workbook(workbook_path)
    write_csv(csv_path, "a,b\n1,2\n")
    write_plain_text(text_path, "repeatable text")
    write_json_value(json_path, {"items": [1, 2, 3]})
    write_docx(docx_path)
    skill = DocumentProcessingSkill()

    executions = (
        ("pdf_reader", PDFReaderInput(path=pdf_path)),
        ("excel_workbook_reader", ExcelWorkbookReaderInput(path=workbook_path)),
        (
            "excel_sheet_reader",
            ExcelSheetReaderInput(path=workbook_path, sheet_name="Summary"),
        ),
        ("csv_reader", CSVReaderInput(path=csv_path, delimiter=",")),
        ("text_reader", TextReaderInput(path=text_path)),
        ("json_reader", JSONReaderInput(path=json_path)),
        ("docx_reader", DOCXReaderInput(path=docx_path)),
    )

    for tool_name, tool_input in executions:
        tool = get_required_tool(skill, tool_name)
        first = tool.execute(tool_input)
        second = tool.execute(tool_input)
        assert first == second


def test_files_remain_readable_after_tool_execution(tmp_path: Path) -> None:
    pdf_path = tmp_path / "document.pdf"
    workbook_path = tmp_path / "workbook.xlsx"
    csv_path = tmp_path / "data.csv"
    text_path = tmp_path / "notes.txt"
    json_path = tmp_path / "document.json"
    docx_path = tmp_path / "document.docx"
    write_text_pdf(pdf_path, ("Resource PDF",))
    write_workbook(workbook_path)
    write_csv(csv_path, "a,b\n1,2\n")
    write_plain_text(text_path, "plain text")
    write_json_value(json_path, {"ok": True})
    write_docx(docx_path)
    skill = DocumentProcessingSkill()

    get_required_tool(skill, "pdf_reader").execute(PDFReaderInput(path=pdf_path))
    get_required_tool(skill, "excel_workbook_reader").execute(
        ExcelWorkbookReaderInput(path=workbook_path)
    )
    get_required_tool(skill, "excel_sheet_reader").execute(
        ExcelSheetReaderInput(path=workbook_path, sheet_name="Summary")
    )
    get_required_tool(skill, "csv_reader").execute(
        CSVReaderInput(path=csv_path, delimiter=",")
    )
    get_required_tool(skill, "text_reader").execute(TextReaderInput(path=text_path))
    get_required_tool(skill, "json_reader").execute(JSONReaderInput(path=json_path))
    get_required_tool(skill, "docx_reader").execute(DOCXReaderInput(path=docx_path))

    assert pdf_path.read_bytes().startswith(b"%PDF")
    workbook = load_workbook(workbook_path, read_only=True)
    try:
        assert workbook.sheetnames == ["Summary", "Formula", "Sparse", "Empty"]
    finally:
        workbook.close()
    assert csv_path.read_text(encoding="utf-8") == "a,b\n1,2\n"
    assert text_path.read_text(encoding="utf-8") == "plain text"
    assert json.loads(json_path.read_text(encoding="utf-8")) == {"ok": True}
    assert Document(docx_path).paragraphs[0].text == "Integration paragraph"
