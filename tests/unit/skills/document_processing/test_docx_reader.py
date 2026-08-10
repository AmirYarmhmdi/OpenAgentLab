import json
from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from openagentlab.skills.document_processing import DocumentProcessingSkill
from openagentlab.skills.document_processing.tools.docx_reader import (
    DOCXReaderError,
    DOCXReaderInput,
    DOCXReaderOutput,
    DOCXReaderTool,
)
from openagentlab.skills.tool import BaseTool


def write_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("")
    document.add_paragraph("  Unicode café سلام  ")
    document.add_heading("Readable heading", level=1)

    first_table = document.add_table(rows=2, cols=3)
    first_table.cell(0, 0).text = "name"
    first_table.cell(0, 1).text = "value"
    first_table.cell(0, 2).text = ""
    first_table.cell(1, 0).text = "alpha"
    first_table.cell(1, 1).text = "42"
    first_table.cell(1, 2).text = "true"

    document.add_paragraph("After first table")

    second_table = document.add_table(rows=1, cols=2)
    second_table.cell(0, 0).text = "شهر"
    second_table.cell(0, 1).text = "München"

    properties = document.core_properties
    properties.title = "DOCX Fixture"
    properties.subject = "Document Processing"
    properties.author = "OpenAgentLab"
    properties.category = "Tests"
    properties.comments = "Small fixture"
    properties.keywords = "docx,reader"
    properties.created = datetime(2026, 8, 8, 9, 30, 0)
    properties.modified = datetime(2026, 8, 8, 10, 45, 0)
    properties.last_modified_by = "Codex"

    document.save(path)


def test_docx_reader_tool_exposes_contract_metadata() -> None:
    tool = DOCXReaderTool()

    assert isinstance(tool, BaseTool)
    assert tool.name == "docx_reader"
    assert tool.description
    assert tool.capability == "document.read.docx"


def test_docx_reader_reads_paragraphs_in_order_with_one_based_indices(
    tmp_path: Path,
) -> None:
    path = tmp_path / "paragraphs.docx"
    write_docx(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))

    assert isinstance(result, DOCXReaderOutput)
    assert result.path == path
    assert result.paragraph_count == 5
    assert [paragraph.index for paragraph in result.paragraphs] == [1, 2, 3, 4, 5]
    assert [paragraph.text for paragraph in result.paragraphs] == [
        "First paragraph",
        "",
        "  Unicode café سلام  ",
        "Readable heading",
        "After first table",
    ]


def test_docx_reader_preserves_empty_unicode_and_whitespace_paragraphs(
    tmp_path: Path,
) -> None:
    path = tmp_path / "paragraph-details.docx"
    write_docx(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))

    assert result.paragraphs[1].text == ""
    assert result.paragraphs[2].text.startswith("  ")
    assert result.paragraphs[2].text.endswith("  ")
    assert "سلام" in result.paragraphs[2].text


def test_docx_reader_reads_tables_in_order_with_one_based_indices(
    tmp_path: Path,
) -> None:
    path = tmp_path / "tables.docx"
    write_docx(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))

    assert result.table_count == 2
    assert [table.index for table in result.tables] == [1, 2]
    assert result.tables[0].rows == (
        ("name", "value", ""),
        ("alpha", "42", "true"),
    )
    assert result.tables[1].rows == (("شهر", "München"),)


def test_docx_reader_preserves_table_values_as_strings_without_header_semantics(
    tmp_path: Path,
) -> None:
    path = tmp_path / "table-values.docx"
    write_docx(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))

    assert result.tables[0].rows[0] == ("name", "value", "")
    assert result.tables[0].rows[1][1] == "42"
    assert isinstance(result.tables[0].rows[1][1], str)
    assert result.tables[0].rows[1][2] == "true"
    assert isinstance(result.tables[0].rows[1][2], str)


def test_docx_reader_extracts_mixed_document_without_interleaved_order_claim(
    tmp_path: Path,
) -> None:
    path = tmp_path / "mixed.docx"
    write_docx(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))

    assert [paragraph.text for paragraph in result.paragraphs] == [
        "First paragraph",
        "",
        "  Unicode café سلام  ",
        "Readable heading",
        "After first table",
    ]
    assert result.tables[0].rows[1] == ("alpha", "42", "true")


def test_docx_reader_returns_basic_metadata_as_serializable_values(
    tmp_path: Path,
) -> None:
    path = tmp_path / "metadata.docx"
    write_docx(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))

    assert result.metadata["title"] == "DOCX Fixture"
    assert result.metadata["subject"] == "Document Processing"
    assert result.metadata["author"] == "OpenAgentLab"
    assert result.metadata["category"] == "Tests"
    assert result.metadata["comments"] == "Small fixture"
    assert result.metadata["keywords"] == "docx,reader"
    assert result.metadata["created"] == "2026-08-08T09:30:00+00:00"
    assert result.metadata["modified"] == "2026-08-08T10:45:00+00:00"
    assert result.metadata["last_modified_by"] == "Codex"


def test_docx_reader_supports_minimal_valid_document(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    Document().save(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))

    assert result.paragraphs == ()
    assert result.tables == ()
    assert result.paragraph_count == 0
    assert result.table_count == 0


def test_docx_reader_rejects_missing_file(tmp_path: Path) -> None:
    path = tmp_path / "missing.docx"

    with pytest.raises(FileNotFoundError, match="missing.docx"):
        DOCXReaderTool().execute(DOCXReaderInput(path=path))


def test_docx_reader_rejects_directory_input(tmp_path: Path) -> None:
    with pytest.raises(IsADirectoryError, match="not a file"):
        DOCXReaderTool().execute(DOCXReaderInput(path=tmp_path))


def test_docx_reader_rejects_unsupported_extension(tmp_path: Path) -> None:
    path = tmp_path / "document.doc"
    path.write_bytes(b"not supported")

    with pytest.raises(ValueError, match=".docx extension"):
        DOCXReaderTool().execute(DOCXReaderInput(path=path))


def test_docx_reader_accepts_uppercase_docx_extension(tmp_path: Path) -> None:
    path = tmp_path / "DOCUMENT.DOCX"
    write_docx(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))

    assert result.paragraph_count == 5


def test_docx_reader_reports_malformed_docx(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a real docx")

    with pytest.raises(DOCXReaderError, match="Could not read DOCX file"):
        DOCXReaderTool().execute(DOCXReaderInput(path=path))


def test_docx_reader_output_serializes_cleanly(tmp_path: Path) -> None:
    path = tmp_path / "serializable.docx"
    write_docx(path)

    result = DOCXReaderTool().execute(DOCXReaderInput(path=path))
    dumped = result.model_dump()
    encoded = json.loads(result.model_dump_json())

    assert dumped["paragraph_count"] == 5
    assert encoded["path"] == str(path)
    assert encoded["paragraphs"][0]["text"] == "First paragraph"
    assert encoded["tables"][0]["rows"][1] == ["alpha", "42", "true"]
    assert "docx.text" not in str(encoded)


def test_docx_reader_is_deterministic(tmp_path: Path) -> None:
    path = tmp_path / "repeatable.docx"
    write_docx(path)
    tool = DOCXReaderTool()
    tool_input = DOCXReaderInput(path=path)

    assert tool.execute(tool_input) == tool.execute(tool_input)


def test_document_processing_skill_exposes_docx_reader_tool() -> None:
    skill = DocumentProcessingSkill()

    assert "document.read.docx" in skill.executable_capabilities
    assert skill.get_tool("docx_reader") is not None
