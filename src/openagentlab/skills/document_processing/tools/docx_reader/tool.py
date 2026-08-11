"""File guide.

- Use: Implements the deterministic tool that reads DOCX files.
- Usage: Import DOCXReaderError, and DOCXReaderTool from
  openagentlab.skills.document_processing.tools.docx_reader.tool.
- Duties: Defines DOCXReaderError, and DOCXReaderTool and related helper logic.
- Depends on: Project modules:
  openagentlab.skills.document_processing.tools.docx_reader.schemas, and
  openagentlab.skills.tool.
"""

from datetime import date, datetime
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from openagentlab.skills.tool import BaseTool

from .schemas import DOCXParagraph, DOCXReaderInput, DOCXReaderOutput, DOCXTable


class DOCXReaderError(ValueError):
    pass


class DOCXReaderTool(BaseTool):
    def __init__(self) -> None:
        super().__init__(
            name="docx_reader",
            description="Extract paragraphs, tables, and metadata from a local .docx.",
            capability="document.read.docx",
            args_schema=DOCXReaderInput,
        )

    def execute(self, tool_input: DOCXReaderInput) -> DOCXReaderOutput:
        path = self._validate_path(tool_input.path)

        try:
            document = Document(path)
        except (BadZipFile, PackageNotFoundError) as exc:
            msg = f"Could not read DOCX file: {path}"
            raise DOCXReaderError(msg) from exc

        paragraphs = tuple(
            DOCXParagraph(index=index, text=paragraph.text)
            for index, paragraph in enumerate(document.paragraphs, start=1)
        )
        tables = tuple(
            DOCXTable(
                index=index,
                rows=tuple(
                    tuple(cell.text for cell in row.cells) for row in table.rows
                ),
            )
            for index, table in enumerate(document.tables, start=1)
        )

        return DOCXReaderOutput(
            path=path,
            paragraphs=paragraphs,
            tables=tables,
            paragraph_count=len(paragraphs),
            table_count=len(tables),
            metadata=self._normalize_metadata(document.core_properties),
        )

    def _validate_path(self, path: Path) -> Path:
        normalized_path = path.expanduser()

        if not normalized_path.exists():
            msg = f"DOCX file does not exist: {normalized_path}"
            raise FileNotFoundError(msg)

        if not normalized_path.is_file():
            msg = f"DOCX input is not a file: {normalized_path}"
            raise IsADirectoryError(msg)

        if normalized_path.suffix.lower() != ".docx":
            msg = f"DOCX input must use a .docx extension: {normalized_path}"
            raise ValueError(msg)

        return normalized_path

    def _normalize_metadata(self, properties: object) -> dict[str, str | None]:
        metadata_fields = (
            "title",
            "subject",
            "author",
            "category",
            "comments",
            "keywords",
            "created",
            "modified",
            "last_modified_by",
        )

        return {
            field: self._serialize_metadata_value(getattr(properties, field, None))
            for field in metadata_fields
        }

    def _serialize_metadata_value(self, value: object) -> str | None:
        if value is None:
            return None

        if isinstance(value, datetime | date):
            return value.isoformat()

        return str(value)
